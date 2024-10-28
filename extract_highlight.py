from docx import Document
import fitz  # PyMuPDF
import os
import sys
import argparse

def extract_highlights_from_pdf(pdf_path):
    """
    Extract highlighted text from a PDF file using PyMuPDF.
    Returns a list of highlighted text snippets.
    """
    # Open the PDF
    doc = fitz.open(pdf_path)
    highlighted_text = []

    # Iterate through all pages
    for page_num in range(len(doc)):
        page = doc[page_num]
        
        # Get all annotations on the page
        annotations = page.annots()
        if annotations:
            for annot in annotations:
                # Check if annotation is a highlight
                if annot.type[0] == 8:  # 8 is the type for highlight annotations
                    # Get the highlighted text
                    rect = annot.rect
                    words = page.get_text("words")
                    highlighted_words = [w[4] for w in words if fitz.Rect(w[:4]).intersects(rect)]
                    text = " ".join(highlighted_words)
                    if text.strip():
                        highlighted_text.append({
                            'page': page_num + 1,
                            'text': text.strip()
                        })
    
    doc.close()
    return highlighted_text

def save_to_word(highlighted_text, output_path):
    """
    Save highlighted text to a Word document.
    """
    doc = Document()
    doc.add_heading('Highlighted Text from PDF', 0)

    current_page = None
    for item in highlighted_text:
        # Add page number as heading if it's a new page
        if current_page != item['page']:
            current_page = item['page']
            doc.add_heading(f'Page {current_page}', level=1)
        
        # Add highlighted text as paragraph
        doc.add_paragraph(item['text'])
    
    doc.save(output_path)

def main():
    # Set up argument parser
    parser = argparse.ArgumentParser(description='Extract highlighted text from a PDF file and save to Word document.')
    parser.add_argument('pdf_path', help='Path to the PDF file')
    parser.add_argument('-o', '--output', help='Output Word document path (optional)')
    
    # Parse arguments
    args = parser.parse_args()
    
    # Check if file exists
    if not os.path.exists(args.pdf_path):
        print(f"Error: File does not exist: {args.pdf_path}")
        return
    
    try:
        # Extract highlights
        print("Extracting highlights...")
        highlights = extract_highlights_from_pdf(args.pdf_path)
        
        if not highlights:
            print("No highlighted text found in the PDF.")
            return
        
        # Create output filename if not specified
        if args.output:
            output_path = args.output
        else:
            output_path = os.path.splitext(args.pdf_path)[0] + "_highlights.docx"
        
        # Save to Word document
        print("Saving to Word document...")
        save_to_word(highlights, output_path)
        
        print(f"Successfully saved highlights to {output_path}")
        print(f"Total highlights found: {len(highlights)}")
        
    except Exception as e:
        print(f"An error occurred: {str(e)}")

if __name__ == "__main__":
    main()